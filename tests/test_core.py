from __future__ import annotations

import json
import tempfile
import unittest
import zipfile
from pathlib import Path
from types import SimpleNamespace

from docx import Document

from reviewer3.core import (
    LLMReviewer,
    FeedbackEntry,
    PrefaceSummary,
    add_comment_to_word_group,
    add_tracked_suggestion,
    default_output_path,
    parse_feedback_entries,
    review_docx,
)
from reviewer3.core import _insert_paragraph_at_start, _preserves_layout_characters, _split_document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class StubReviewer:
    def __init__(self) -> None:
        self.review_calls: list[dict] = []
        self.preface_calls: list[dict] = []

    def review_document(self, document_text):
        self.review_calls.append({"document_text": document_text})
        return [
            FeedbackEntry(
                feedback_type="local_modification",
                word_group="abbreviations",
                comment=None,
                suggestion="abbreviations (LLM)",
            ),
            FeedbackEntry(
                feedback_type="regional_question",
                word_group="explains the method",
                comment="Could you provide one concrete example of the method output?",
                suggestion=None,
            ),
            FeedbackEntry(
                feedback_type="global_feedback",
                word_group=None,
                comment="Some method details and interpretation appear in different sections and should be connected explicitly.",
                suggestion=None,
            ),
        ]

    def compose_preface(self, document_text, feedback_entries):
        self.preface_calls.append(
            {
                "document_text": document_text,
                "feedback_entries": list(feedback_entries),
            }
        )
        return PrefaceSummary(
            about_document="This document introduces the study goals and outlines the applied method.",
            scientific_importance="The topic is valuable because it improves methodological clarity for reproducible research.",
            global_feedback="Open questions remain around how results are interpreted across sections, and related claims should be grouped more tightly.",
            feedback_summary="Local wording and punctuation were improved, and one regional clarification question was raised.",
        )


class _StreamingCompletionsStub:
    def __init__(self, chunks: list[str]) -> None:
        self._chunks = chunks
        self.calls: list[dict] = []

    def create(self, **kwargs):
        self.calls.append(kwargs)
        return [SimpleNamespace(choices=[SimpleNamespace(delta=SimpleNamespace(content=chunk))]) for chunk in self._chunks]


class _StreamingClientStub:
    def __init__(self, chunks: list[str]) -> None:
        self.completions = _StreamingCompletionsStub(chunks)
        self.chat = SimpleNamespace(completions=self.completions)


class TestCore(unittest.TestCase):
    def test_insert_paragraph_at_start_formats_string_list_as_bullets(self):
        doc = Document()
        doc.add_paragraph("Existing text")

        paragraph = _insert_paragraph_at_start(doc, ["First item", "Second item"])

        self.assertEqual("• First item\n• Second item", paragraph.text)

    def test_insert_paragraph_at_start_keeps_string_unchanged(self):
        doc = Document()

        paragraph = _insert_paragraph_at_start(doc, "Plain text")

        self.assertEqual("Plain text", paragraph.text)

    def test_split_document_avoids_cutting_inside_brackets_when_possible(self):
        text = "Prefix words (a bracketed phrase) suffix words."
        chunks = _split_document(text, 35)

        self.assertEqual(text, "".join(chunks))
        self.assertTrue(all(len(chunk) <= 35 for chunk in chunks))
        self.assertFalse(any("(" in chunk and ")" not in chunk for chunk in chunks))

    def test_layout_guard_rejects_line_break_or_bracket_changes(self):
        self.assertTrue(_preserves_layout_characters("alpha (beta)\ngamma", "delta (beta)\nepsilon"))
        self.assertFalse(_preserves_layout_characters("alpha (beta)", "alpha beta"))
        self.assertFalse(_preserves_layout_characters("alpha\nbeta", "alpha beta"))

    def test_regional_comment_is_anchored_to_word_group(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Text before ")
        emphasized = paragraph.add_run("the word group")
        emphasized.bold = True
        paragraph.add_run(" and text after.")

        add_comment_to_word_group(doc, paragraph, "the word group", "Could you clarify this?")

        xml = paragraph._p.xml
        range_start = xml.index("<w:commentRangeStart")
        target = xml.index("the word group")
        range_end = xml.index("<w:commentRangeEnd")
        self.assertLess(xml.index("Text before "), range_start)
        self.assertLess(range_start, target)
        self.assertLess(target, range_end)
        self.assertLess(range_end, xml.index(" and text after."))
        self.assertTrue(emphasized.bold)

    def test_tracked_suggestion_preserves_zotero_and_footnote_xml(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Improve this wording")

        zotero_field = OxmlElement("w:fldSimple")
        zotero_field.set(qn("w:instr"), "ADDIN ZOTERO_ITEM CSL_CITATION")
        paragraph._p.append(zotero_field)
        footnote_run = OxmlElement("w:r")
        footnote_reference = OxmlElement("w:footnoteReference")
        footnote_reference.set(qn("w:id"), "1")
        footnote_run.append(footnote_reference)
        paragraph._p.append(footnote_run)

        next_id = add_tracked_suggestion(paragraph, "Clarify this wording", change_id=7)

        xml = paragraph._p.xml
        self.assertEqual(9, next_id)
        self.assertIn("ZOTERO_ITEM", xml)
        self.assertIn("footnoteReference", xml)
        self.assertIn("w:del", xml)
        self.assertIn("w:ins", xml)

    def test_default_output_path(self):
        self.assertEqual(
            default_output_path(Path("/tmp/paper.docx")),
            Path("/tmp/paper_rv3.docx"),
        )

    def test_review_docx_adds_structured_feedback_and_preface(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "paper.docx"
            output_path = Path(tmpdir) / "paper_rv3.docx"
            doc = Document()
            doc.add_paragraph("Paragraph one about abbreviations.")
            doc.add_paragraph("Paragraph two explains the method.")
            doc.save(str(input_path))

            stub = StubReviewer()
            saved = review_docx(input_docx=input_path, output_docx=output_path, reviewer=stub)

            self.assertEqual(saved, output_path)
            self.assertTrue(output_path.exists())
            self.assertTrue(input_path.exists())
            self.assertEqual(1, len(stub.review_calls))
            self.assertEqual(1, len(stub.preface_calls))
            self.assertIn("Paragraph one about abbreviations.", stub.review_calls[0]["document_text"])
            self.assertIn("Paragraph two explains the method.", stub.review_calls[0]["document_text"])

            reviewed = Document(str(output_path))
            self.assertEqual("Reviewer3 Preface", reviewed.paragraphs[0].text)
            self.assertIn("What this document is about:", reviewed.paragraphs[1].text)
            self.assertIn("Scientific importance:", reviewed.paragraphs[2].text)
            self.assertIn("Global feedback:", reviewed.paragraphs[3].text)
            self.assertIn("Short feedback summary:", reviewed.paragraphs[4].text)

            comments = list(reviewed.comments)
            self.assertEqual(1, len(comments))
            self.assertIn("concrete example", comments[0].text)

            with zipfile.ZipFile(output_path) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
                settings_xml = archive.read("word/settings.xml").decode("utf-8")
            self.assertIn("<w:trackRevisions", settings_xml)
            self.assertIn("<w:pPr><w:rPr><w:ins", document_xml)
            method_paragraph = document_xml[document_xml.index("Paragraph two") :]
            method_paragraph = method_paragraph[: method_paragraph.index("</w:p>")]
            comment_start = method_paragraph.index("<w:commentRangeStart")
            comment_end = method_paragraph.index("<w:commentRangeEnd")
            self.assertLess(method_paragraph.index("Paragraph two"), comment_start)
            self.assertLess(comment_start, method_paragraph.index("explains the method"))
            self.assertLess(method_paragraph.index("explains the method"), comment_end)

            xml = reviewed.part.element.xml
            self.assertIn("w:ins", xml)
            self.assertIn("[reviewer3 suggestion]", xml)
            self.assertIn("abbreviations (LLM)", xml)
            self.assertIn("Open questions remain", xml)

    def test_review_docx_avoids_overwriting_existing_output(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / "paper.docx"
            requested_output = Path(tmpdir) / "paper_rv3.docx"
            existing_1 = Path(tmpdir) / "paper_rv3_1.docx"
            expected_output = Path(tmpdir) / "paper_rv3_2.docx"

            input_doc = Document()
            input_doc.add_paragraph("Paragraph one about abbreviations.")
            input_doc.save(str(input_path))

            existing_doc = Document()
            existing_doc.add_paragraph("Do not overwrite this file.")
            existing_doc.save(str(requested_output))
            existing_doc.save(str(existing_1))

            stub = StubReviewer()
            saved = review_docx(input_docx=input_path, output_docx=requested_output, reviewer=stub)

            self.assertEqual(expected_output, saved)
            self.assertTrue(requested_output.exists())
            self.assertTrue(existing_1.exists())
            self.assertTrue(expected_output.exists())
            self.assertEqual("Do not overwrite this file.", Document(str(requested_output)).paragraphs[0].text)

    def test_parse_feedback_entries_accepts_jsonl(self):
        raw = "\n".join(
            [
                '{"feedback_type":"local_modification","word_group":"alpha","comment":null,"suggestion":"beta"}',
                '{"feedback_type":"regional_question","word_group":"gamma","comment":"Could you clarify this point?","suggestion":null}',
                '{"feedback_type":"global_feedback","word_group":null,"comment":"Please improve section transitions.","suggestion":null}',
            ]
        )

        entries = parse_feedback_entries(raw)
        self.assertEqual(3, len(entries))
        self.assertEqual("local_modification", entries[0].feedback_type)
        self.assertEqual("alpha", entries[0].word_group)
        self.assertEqual("regional_question", entries[1].feedback_type)
        self.assertEqual("Could you clarify this point?", entries[1].comment)
        self.assertEqual("global_feedback", entries[2].feedback_type)

    def test_review_document_streams_jsonl_and_reports_position_progress(self):
        chunks = [
            '{"feedback_type":"local_modification","word_group":"Beta term","comment":null,"suggestion":"Better term"}\n',
            '{"feedback_type":"regional_question","word_group":"Gamma section","comment":"Could you add one concrete example?","suggestion":null}\n',
        ]
        client = _StreamingClientStub(chunks)
        progress_messages: list[str] = []
        reviewer = LLMReviewer(
            client=client,
            model="test-model",
            progress_callback=progress_messages.append,
        )

        entries = reviewer.review_document("Alpha text. Beta term appears here. Gamma section appears later.")

        self.assertEqual(2, len(entries))
        self.assertEqual("local_modification", entries[0].feedback_type)
        self.assertEqual("regional_question", entries[1].feedback_type)
        self.assertGreaterEqual(len(progress_messages), 2)
        self.assertTrue(any("document position" in message for message in progress_messages))
        self.assertEqual(True, client.completions.calls[0]["stream"])

    def test_review_document_splits_long_text_into_bounded_requests(self):
        chunks = [
            '{"feedback_type":"local_modification","word_group":"Alpha","comment":null,"suggestion":"Beta"}\n'
        ]
        client = _StreamingClientStub(chunks)
        reviewer = LLMReviewer(client=client, model="test-model", chunk_size=30)

        reviewer.review_document(("Alpha short paragraph.\n\n" * 5).strip())

        self.assertGreater(len(client.completions.calls), 1)
        for call in client.completions.calls:
            payload = call["messages"][1]["content"]
            self.assertLessEqual(len(json.loads(payload)["document"]), 30)


if __name__ == "__main__":
    unittest.main()
