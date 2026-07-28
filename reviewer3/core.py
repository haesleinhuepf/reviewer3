from __future__ import annotations

import json
import os
import re
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from openai import OpenAI, __version__
import yaml

REVIEWER_AUTHOR = "reviewer3"
REVIEWER_INITIALS = "rv3"
DEFAULT_CHUNK_SIZE = 12_000


@dataclass
class FeedbackEntry:
    feedback_type: str
    word_group: str | None
    comment: str | None
    suggestion: str | None


@dataclass
class PrefaceSummary:
    about_document: str
    scientific_importance: str
    global_feedback: str
    feedback_summary: str


class LLMReviewer:
    def __init__(
        self,
        client: OpenAI,
        model: str,
        temperature: float = 0.2,
        progress_callback: Callable[[str], None] | None = None,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
    ) -> None:
        self._client = client
        self._model = model
        self._temperature = temperature
        self._progress_callback = progress_callback
        if chunk_size < 1:
            raise ValueError("chunk_size must be positive")
        self._chunk_size = chunk_size

    @classmethod
    def from_env(cls, progress_callback: Callable[[str], None] | None = None) -> "LLMReviewer":
        base_url = os.getenv("REVIEWER3_BASE_URL", "http://localhost:11434/v1")
        api_key = os.getenv("REVIEWER3_API_KEY", "ollama")
        model = os.getenv("REVIEWER3_MODEL", "gpt-oss:20b")
        temperature = float(os.getenv("REVIEWER3_TEMPERATURE", "0.2"))
        chunk_size = int(os.getenv("REVIEWER3_CHUNK_SIZE", str(DEFAULT_CHUNK_SIZE)))
        return cls(
            client=OpenAI(base_url=base_url, api_key=api_key),
            model=model,
            temperature=temperature,
            progress_callback=progress_callback,
            chunk_size=chunk_size,
        )

    def review_document(self, document_text: str) -> list[FeedbackEntry]:
        chunks = _split_document(document_text, self._chunk_size)
        entries: list[FeedbackEntry] = []
        for index, chunk in enumerate(chunks, start=1):
            if self._progress_callback and len(chunks) > 1:
                self._progress_callback(f"Reviewing document chunk {index}/{len(chunks)}...")
            entries.extend(self._review_chunk(chunk, index, len(chunks)))
        return _deduplicate_feedback(entries)

    def _review_chunk(self, document_text: str, chunk_number: int, chunk_count: int) -> list[FeedbackEntry]:
        messages = [
            {
                "role": "system",
                "content": (
                    "You are reviewer3, a constructive reviewer for scientific writing. "
                    "Return JSONL only. "
                    "Review the supplied document text and return one feedback entry per line as a JSON object. "
                    "Each JSON line must include: feedback_type, word_group, comment, suggestion. "
                    "Allowed feedback_type values are: local_modification, regional_question, global_feedback. "
                    "local_modification: use word_group and suggestion (small edits only, such as wording, punctuation, or rephrasing). "
                    "regional_question: use word_group and comment as a constructive question. "
                    "global_feedback: document-level feedback with comment, word_group may be null, suggestion must be null. "
                    "Do not rewrite entire sections if a local edit is enough. "
                    "Preserve all line breaks and every bracket character in a local replacement. "
                    "Never add or remove (), [], {}, or <> merely because the supplied text starts or ends mid-paragraph. "
                    "Do not propose a local modification whose word_group crosses a line break. "
                    "Use concise, actionable, and respectful feedback. "
                    "Write every comment in a natural first-person reviewer voice, as an experienced senior scientist. "
                    "Avoid templated boilerplate such as 'I read this and now I think ...'. "
                    "Keep this tone friendly, professional, and constructive. "
                    "Do not wrap output in markdown fences and do not emit a list or outer object."
                ),
            },
            {
                "role": "user",
                "content": json.dumps(
                    {
                        "document": document_text,
                        "document_chunk": {"number": chunk_number, "total": chunk_count},
                        "instructions": {
                            "output_format": "JSONL (one JSON object per line)",
                            "line_schema": {
                                "feedback_type": "local_modification | regional_question | global_feedback",
                                "word_group": "exact text span (no '...' in it!) this feedback refers to, or null for global_feedback",
                                "comment": "question/comment text or null",
                                "suggestion": "replacement for word_group or null"
                            }
                        },
                    },
                    ensure_ascii=False,
                ),
            },
        ]
        try:
            stream = self._client.chat.completions.create(
                model=self._model,
                temperature=self._temperature,
                messages=messages,
                stream=True,
                tools=[],
                tool_choice="none",
            )
            return _consume_feedback_stream(
                stream=stream,
                document_text=document_text,
                progress_callback=self._progress_callback,
            )
        except Exception as exc:
            if self._progress_callback:
                self._progress_callback(f"Review error: {type(exc).__name__}: {exc}")
            raise

    def compose_preface(self, document_text: str, feedback_entries: list[FeedbackEntry]) -> PrefaceSummary:
        global_feedback = [entry.comment for entry in feedback_entries if entry.feedback_type == "global_feedback" and entry.comment]
        local_count = sum(1 for entry in feedback_entries if entry.feedback_type == "local_modification")
        regional_count = sum(1 for entry in feedback_entries if entry.feedback_type == "regional_question")

        messages = [
            {
                "role": "system",
                "content": (
                    "You are reviewer3, a friendly and constructive reviewer for scientific writing. "
                    "Create a concise preface with exactly the requested sections. "
                    "Return JSON only with keys: about_document, scientific_importance, global_feedback, feedback_summary. "
                    "Write 2-3 sentences for about_document and 2-3 sentences for scientific_importance. "
                    "Use global_feedback for open questions and cross-section coherence points. "
                    "Use feedback_summary for a very short summary of local and regional feedback. "
                    "End feedback_summary with a positive sentence expressing confidence that, once the raised points are addressed, "
                    "the manuscript will make an excellent contribution. "
                    "Write every section in a natural first-person reviewer voice, as an experienced senior scientist. "
                    "Avoid templated boilerplate such as 'I read this and now I think ...'. "
                    "Keep this tone friendly, professional, and constructive. Motivate the author in a positive way. "
                ),
            },
            {
                "role": "user",
                "content": json.dumps(
                    {
                        "document": _document_overview(document_text, self._chunk_size),
                        "global_feedback_entries": global_feedback,
                        "local_modification_count": local_count,
                        "regional_question_count": regional_count,
                        "instructions": {
                            "about_document": "What the document is about in 2-3 sentences.",
                            "scientific_importance": "Why it is outstanding/important for science or the scientific community in 2-3 sentences.",
                            "global_feedback": "General open questions and points that should be connected across sections.",
                            "feedback_summary": "Very short summary of collected local and regional feedback.",
                        },
                    },
                    ensure_ascii=False,
                ),
            },
        ]
        try:
            response = self._client.chat.completions.create(
                model=self._model,
                temperature=self._temperature,
                messages=messages,
            )
            content = response.choices[0].message.content or "{}"
            return parse_preface_summary(
                content,
                fallback_global_feedback=global_feedback,
                fallback_local_count=local_count,
                fallback_regional_count=regional_count,
            )
        except Exception as exc:
            if self._progress_callback:
                self._progress_callback(f"Preface error: {type(exc).__name__}: {exc}")
            raise


def parse_preface_summary(
    raw: str,
    fallback_global_feedback: list[str],
    fallback_local_count: int,
    fallback_regional_count: int,
) -> PrefaceSummary:
    parsed = _parse_json_object(raw)
    about_document = _to_optional_str(parsed.get("about_document"))
    scientific_importance = _to_optional_str(parsed.get("scientific_importance"))
    global_feedback = _to_optional_str(parsed.get("global_feedback"))
    feedback_summary = _to_optional_str(parsed.get("feedback_summary"))

    if not about_document:
        about_document = _fallback_preface_summary(
            fallback_global_feedback=fallback_global_feedback,
            fallback_local_count=fallback_local_count,
            fallback_regional_count=fallback_regional_count,
        ).about_document
    if not scientific_importance:
        scientific_importance = _fallback_preface_summary(
            fallback_global_feedback=fallback_global_feedback,
            fallback_local_count=fallback_local_count,
            fallback_regional_count=fallback_regional_count,
        ).scientific_importance
    if not global_feedback:
        global_feedback = _fallback_preface_summary(
            fallback_global_feedback=fallback_global_feedback,
            fallback_local_count=fallback_local_count,
            fallback_regional_count=fallback_regional_count,
        ).global_feedback
    if not feedback_summary:
        feedback_summary = _fallback_preface_summary(
            fallback_global_feedback=fallback_global_feedback,
            fallback_local_count=fallback_local_count,
            fallback_regional_count=fallback_regional_count,
        ).feedback_summary

    base_url = os.getenv("REVIEWER3_BASE_URL", "http://localhost:11434/v1")
    model = os.getenv("REVIEWER3_MODEL", "gpt-oss:20b")

    feedback_summary = feedback_summary + f"""

This review was AI-generated using reviewer3 tool (version {__version__}) with the model {model} at {base_url}.
Read more: https://github.com/haesleinhuepf/reviewer3
"""
    
    return PrefaceSummary(
        about_document=about_document,
        scientific_importance=scientific_importance,
        global_feedback=global_feedback,
        feedback_summary=feedback_summary,
    )


def parse_feedback_entries(raw: str) -> list[FeedbackEntry]:
    entries: list[FeedbackEntry] = []
    for item in _parse_json_lines(raw):
        entry = _entry_from_dict(item)
        if entry:
            entries.append(entry)

    if entries:
        return entries

    # Backward compatibility for non-JSONL responses.
    parsed = _parse_json_array_or_wrapped(raw)
    for item in parsed:
        entry = _entry_from_dict(item)
        if entry:
            entries.append(entry)
    return entries


def _fallback_preface_summary(
    fallback_global_feedback: list[str],
    fallback_local_count: int,
    fallback_regional_count: int,
) -> PrefaceSummary:
    return PrefaceSummary(
        about_document="I think the document presents a clear scientific contribution and outlines its methods and findings.",
        scientific_importance=(
            "I think the work is relevant to the scientific community because it clarifies the research context and provides reproducible insights."
        ),
        global_feedback=(
            " ".join(fallback_global_feedback)
            if fallback_global_feedback
            else "I think there are no major global issues, but the narrative can be strengthened by linking related concepts across sections."
        ),
        feedback_summary=(
            "I think the collected feedback is focused and actionable: "
            f"{fallback_local_count} local modifications and {fallback_regional_count} regional questions. "
            "Once these points are addressed, I am confident that the manuscript will make an excellent contribution."
        ),
    )


def _consume_feedback_stream(
    stream: Any,
    document_text: str,
    progress_callback: Callable[[str], None] | None,
) -> list[FeedbackEntry]:
    entries: list[FeedbackEntry] = []
    full_text = ""
    text_buffer = ""
    normalized_text = document_text.lower()
    text_length = max(1, len(document_text))

    def report_progress(word_group: str | None, parsed_count: int) -> None:
        if not progress_callback:
            return
        if not word_group:
            progress_callback(f"Parsed {parsed_count} feedback entries...")
            return
        idx = normalized_text.find(word_group.lower())
        if idx == -1:
            progress_callback(f"Parsed {parsed_count} feedback entries (word_group not found yet).")
            return
        pct = int((idx / text_length) * 100)
        progress_callback(
            f"Parsed {parsed_count} feedback entries... document position {pct}% ({idx + 1}/{len(document_text)})"
        )

    for chunk in stream:
        delta = chunk.choices[0].delta.content if chunk.choices else None
        if not delta:
            continue
        full_text += delta
        text_buffer += delta
        lines = text_buffer.splitlines()
        if text_buffer and not text_buffer.endswith("\n"):
            text_buffer = lines.pop() if lines else text_buffer
        else:
            text_buffer = ""

        for line in lines:
            parsed_line = _parse_json_line(line)
            if not parsed_line:
                continue
            entry = _entry_from_dict(parsed_line)
            if not entry:
                continue
            entries.append(entry)
            report_progress(entry.word_group, len(entries))

    trailing_line = text_buffer.strip()
    if trailing_line:
        parsed_line = _parse_json_line(trailing_line)
        if parsed_line:
            entry = _entry_from_dict(parsed_line)
            if entry:
                entries.append(entry)
                report_progress(entry.word_group, len(entries))

    if not entries:
        entries = parse_feedback_entries(full_text)
    return entries


def _split_document(document_text: str, max_chars: int) -> list[str]:
    """Split at paragraph boundaries and avoid cutting through words or brackets."""
    if not document_text:
        return []

    chunks: list[str] = []
    current = ""
    for paragraph in document_text.split("\n\n"):
        parts = _split_long_paragraph(paragraph, max_chars)
        for part in parts:
            separator = "\n\n" if current else ""
            if current and len(current) + len(separator) + len(part) > max_chars:
                chunks.append(current)
                current = part
            else:
                current += separator + part
    if current:
        chunks.append(current)
    return chunks


def _split_long_paragraph(paragraph: str, max_chars: int) -> list[str]:
    if not paragraph:
        return [""]

    parts: list[str] = []
    start = 0
    bracket_pairs = {")": "(", "]": "[", "}": "{", ">": "<"}
    while len(paragraph) - start > max_chars:
        limit = start + max_chars
        stack: list[str] = []
        safe_breaks: list[int] = []
        for index in range(start, limit):
            char = paragraph[index]
            if char in "([{<":
                stack.append(char)
            elif char in bracket_pairs and stack and stack[-1] == bracket_pairs[char]:
                stack.pop()
            if not stack and (char.isspace() or char in ".!?;:"):
                safe_breaks.append(index + 1)

        end = safe_breaks[-1] if safe_breaks else limit
        parts.append(paragraph[start:end])
        start = end
    parts.append(paragraph[start:])
    return parts


def _document_overview(document_text: str, max_chars: int) -> str:
    """Create a bounded, representative sample for the document-level preface."""
    if len(document_text) <= max_chars:
        return document_text
    separator = "\n\n[...]\n\n"
    sample_size = max(1, (max_chars - 2 * len(separator)) // 3)
    middle_start = max(0, (len(document_text) - sample_size) // 2)
    return separator.join(
        (
            document_text[:sample_size],
            document_text[middle_start : middle_start + sample_size],
            document_text[-sample_size:],
        )
    )[:max_chars]


def _deduplicate_feedback(entries: list[FeedbackEntry]) -> list[FeedbackEntry]:
    unique: list[FeedbackEntry] = []
    seen: set[tuple[str, str | None, str | None, str | None]] = set()
    for entry in entries:
        key = (entry.feedback_type, entry.word_group, entry.comment, entry.suggestion)
        if key not in seen:
            seen.add(key)
            unique.append(entry)
    return unique


def _parse_json_lines(raw: str) -> list[dict[str, Any]]:
    parsed: list[dict[str, Any]] = []
    for line in raw.splitlines():
        item = _parse_json_line(line)
        if item:
            parsed.append(item)
    return parsed


def _parse_json_line(raw_line: str) -> dict[str, Any] | None:
    line = raw_line.strip()
    if not line:
        return None
    if line.startswith("```"):
        return None
    try:
        item = json.loads(line)
    except json.JSONDecodeError:
        return None
    return item if isinstance(item, dict) else None


def _entry_from_dict(item: dict[str, Any]) -> FeedbackEntry | None:
    feedback_type = _normalize_feedback_type(item.get("feedback_type") or item.get("type") or item.get("kind"))
    word_group = _to_optional_str(item.get("word_group") or item.get("text_span") or item.get("target"))
    comment = _to_optional_str(item.get("comment") or item.get("question"))
    suggestion = _to_optional_str(item.get("suggestion") or item.get("replacement") or item.get("change_to"))

    if feedback_type == "local_modification":
        if not word_group or not suggestion:
            return None
        return FeedbackEntry(
            feedback_type=feedback_type,
            word_group=word_group,
            comment=comment,
            suggestion=suggestion,
        )

    if feedback_type == "regional_question":
        if not word_group or not comment:
            return None
        return FeedbackEntry(
            feedback_type=feedback_type,
            word_group=word_group,
            comment=comment,
            suggestion=None,
        )

    if feedback_type == "global_feedback":
        if not comment:
            comment = suggestion
        if not comment:
            return None
        return FeedbackEntry(
            feedback_type=feedback_type,
            word_group=None,
            comment=comment,
            suggestion=None,
        )

    return None


def _parse_json_object(raw: str) -> dict[str, Any]:
    text = raw.strip()
    if not text:
        return {}
    try:
        parsed = json.loads(text)
        return parsed if isinstance(parsed, dict) else {}
    except json.JSONDecodeError:
        pass
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return {}
    try:
        parsed = json.loads(text[start : end + 1])
        return parsed if isinstance(parsed, dict) else {}
    except json.JSONDecodeError:
        return {}


def _parse_json_array_or_wrapped(raw: str) -> list[dict[str, Any]]:
    text = raw.strip()
    if not text:
        return []

    candidates: list[Any] = []
    try:
        candidates.append(json.loads(text))
    except json.JSONDecodeError:
        pass

    for start_char, end_char in (("[", "]"), ("{", "}")):
        start = text.find(start_char)
        end = text.rfind(end_char)
        if start == -1 or end == -1 or end <= start:
            continue
        snippet = text[start : end + 1]
        try:
            candidates.append(json.loads(snippet))
        except json.JSONDecodeError:
            continue

    for candidate in candidates:
        if isinstance(candidate, dict) and isinstance(candidate.get("feedback_entries"), list):
            return [item for item in candidate["feedback_entries"] if isinstance(item, dict)]
        if isinstance(candidate, list):
            return [item for item in candidate if isinstance(item, dict)]
    return []


def _to_str_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    text = str(value).strip()
    return [text] if text else []


def _to_optional_str(value: Any) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    return text or None


def _normalize_feedback_type(value: Any) -> str:
    text = str(value or "").strip().lower()
    mapping = {
        "local_modification": "local_modification",
        "local": "local_modification",
        "regional_question": "regional_question",
        "regional": "regional_question",
        "question": "regional_question",
        "global_feedback": "global_feedback",
        "global": "global_feedback",
    }
    return mapping.get(text, "")


def default_output_path(input_docx: Path) -> Path:
    return input_docx.with_name(f"{input_docx.stem}_rv3.docx")


def _non_overwriting_output_path(path: Path) -> Path:
    candidate = Path(path)
    if not candidate.exists():
        return candidate

    suffix = "".join(candidate.suffixes)
    if suffix:
        base_name = candidate.name[: -len(suffix)]
    else:
        base_name = candidate.name

    counter = 1
    while True:
        numbered = candidate.with_name(f"{base_name}_{counter}{suffix}")
        if not numbered.exists():
            return numbered
        counter += 1


def add_comment_to_paragraph(document: Document, paragraph: Any, comment_text: str) -> None:
    if not comment_text.strip():
        return
    if not paragraph.runs:
        paragraph.add_run(" ")
    document.add_comment(paragraph.runs, text=comment_text, author=REVIEWER_AUTHOR, initials=REVIEWER_INITIALS)


def add_comment_to_word_group(
    document: Document,
    paragraph: Any,
    word_group: str,
    comment_text: str,
) -> None:
    """Add a comment whose range is the first matching word group."""
    if not word_group.strip() or not comment_text.strip():
        return

    runs = list(paragraph.runs)
    run_text = "".join(run.text for run in runs)
    match = re.search(re.escape(word_group.strip()), run_text, flags=re.IGNORECASE)
    if match is None:
        return

    selected_runs: list[Run] = []
    offset = 0
    for run in runs:
        run_end = offset + len(run.text)
        selection_start = max(match.start(), offset)
        selection_end = min(match.end(), run_end)
        if selection_start < selection_end:
            selected_runs.append(
                _isolate_run_text(
                    run,
                    selection_start - offset,
                    selection_end - offset,
                )
            )
        offset = run_end

    if selected_runs:
        document.add_comment(
            selected_runs,
            text=comment_text,
            author=REVIEWER_AUTHOR,
            initials=REVIEWER_INITIALS,
        )


def _isolate_run_text(run: Run, start: int, end: int) -> Run:
    """Split a run and return the run containing ``text[start:end]``."""
    if start == 0 and end == len(run.text):
        return run

    original_text = run.text
    fragments = (
        (original_text[:start], False),
        (original_text[start:end], True),
        (original_text[end:], False),
    )
    selected_run: Run | None = None
    for text, is_selected in fragments:
        if not text:
            continue
        run_xml = deepcopy(run._r)
        fragment = Run(run_xml, run._parent)
        fragment.text = text
        run._r.addprevious(run_xml)
        if is_selected:
            selected_run = fragment

    run._r.getparent().remove(run._r)
    assert selected_run is not None
    return selected_run


def add_tracked_suggestion(paragraph: Any, word_group: str, suggestion: str, change_id: int) -> int:
    if not word_group.strip() or not suggestion:
        return change_id

    runs = list(paragraph.runs)
    original = "".join(run.text for run in runs)
    if not original:
        return change_id

    match = re.search(re.escape(word_group), original, flags=re.IGNORECASE)
    if match is None:
        if original.strip().lower() != word_group.strip().lower():
            return change_id
        match_start = 0
        match_end = len(original)
    else:
        match_start = match.start()
        match_end = match.end()

    matched_text = original[match_start:match_end]
    if matched_text == suggestion:
        return change_id

    common_prefix_length = 0
    for original_char, suggestion_char in zip(matched_text, suggestion):
        if original_char != suggestion_char:
            break
        common_prefix_length += 1

    common_suffix_length = 0
    max_suffix_length = min(
        len(matched_text) - common_prefix_length,
        len(suggestion) - common_prefix_length,
    )
    while (
        common_suffix_length < max_suffix_length
        and matched_text[-common_suffix_length - 1] == suggestion[-common_suffix_length - 1]
    ):
        common_suffix_length += 1

    match_start += common_prefix_length
    match_end -= common_suffix_length
    suggestion_end = len(suggestion) - common_suffix_length if common_suffix_length else len(suggestion)
    inserted_text = suggestion[common_prefix_length:suggestion_end]

    offset = 0
    for run in runs:
        run_end = offset + len(run.text)
        if offset <= match_start and match_end <= run_end:
            if _replace_text_in_simple_run(
                run,
                match_start - offset,
                match_end - offset,
                inserted_text,
                change_id,
            ):
                deleted_text = run.text[match_start - offset : match_end - offset]
                return change_id + bool(deleted_text) + bool(inserted_text)
            return change_id
        offset = run_end

    # A change spanning runs or an unsupported inline object is deliberately
    # skipped rather than risking damage to document structure.
    return change_id


def _replace_text_in_simple_run(
    run: Any,
    start: int,
    end: int,
    inserted_text: str,
    change_id: int,
) -> bool:
    run_element = run._r
    children = list(run_element)
    text_nodes = [child for child in children if child.tag == qn("w:t")]
    if len(text_nodes) != 1 or any(child.tag not in {qn("w:rPr"), qn("w:t")} for child in children):
        return False

    run_text = text_nodes[0].text or ""
    if not (0 <= start <= end <= len(run_text)):
        return False
    deleted_text = run_text[start:end]
    if not deleted_text and not inserted_text:
        return False

    parent = run_element.getparent()
    index = parent.index(run_element)
    run_properties = run_element.find(qn("w:rPr"))
    replacement_nodes: list[Any] = []

    before = run_text[:start]
    after = run_text[end:]
    if before:
        replacement_nodes.append(_make_text_run(before, run_properties))
    if deleted_text:
        replacement_nodes.append(_make_tracked_change("w:del", "w:delText", deleted_text, change_id, run_properties))
        change_id += 1
    if inserted_text:
        replacement_nodes.append(_make_tracked_change("w:ins", "w:t", inserted_text, change_id, run_properties))
    if after:
        replacement_nodes.append(_make_text_run(after, run_properties))

    parent.remove(run_element)
    for node in reversed(replacement_nodes):
        parent.insert(index, node)
    return True


def _make_text_run(text_value: str, run_properties: Any | None) -> Any:
    run = OxmlElement("w:r")
    if run_properties is not None:
        run.append(deepcopy(run_properties))
    text = OxmlElement("w:t")
    if text_value[:1].isspace() or text_value[-1:].isspace():
        text.set(qn("xml:space"), "preserve")
    text.text = text_value
    run.append(text)
    return run


def _make_tracked_change(
    change_tag: str,
    text_tag: str,
    text_value: str,
    change_id: int,
    run_properties: Any | None,
) -> Any:
    change = OxmlElement(change_tag)
    change.set(qn("w:id"), str(change_id))
    change.set(qn("w:author"), REVIEWER_AUTHOR)
    change.set(qn("w:date"), datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z"))
    run = _make_text_run(text_value, run_properties)
    run[-1].tag = qn(text_tag)
    change.append(run)
    return change


def _append_tracked_insert(paragraph: Any, inserted_text: str, change_id: int) -> None:
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(change_id))
    ins.set(qn("w:author"), REVIEWER_AUTHOR)
    ins.set(qn("w:date"), datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z"))

    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    if inserted_text[:1].isspace() or inserted_text[-1:].isspace():
        text.set(qn("xml:space"), "preserve")
    text.text = inserted_text
    run.append(text)
    ins.append(run)

    paragraph._p.append(ins)


def _append_tracked_delete(paragraph: Any, deleted_text: str, change_id: int) -> None:
    deletion = OxmlElement("w:del")
    deletion.set(qn("w:id"), str(change_id))
    deletion.set(qn("w:author"), REVIEWER_AUTHOR)
    deletion.set(qn("w:date"), datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z"))

    run = OxmlElement("w:r")
    text = OxmlElement("w:delText")
    if deleted_text[:1].isspace() or deleted_text[-1:].isspace():
        text.set(qn("xml:space"), "preserve")
    text.text = deleted_text
    run.append(text)
    deletion.append(run)

    paragraph._p.append(deletion)


def _next_change_id(document: Document) -> int:
    max_id = 0
    body = document._body._element
    for tag in (qn("w:ins"), qn("w:del")):
        for node in body.iter(tag):
            raw_id = node.get(qn("w:id"))
            if raw_id is None:
                continue
            try:
                value = int(raw_id)
            except ValueError:
                continue
            if value > max_id:
                max_id = value
    return max_id + 1


def _enable_track_revisions(document: Document) -> None:
    settings = document.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        settings.append(OxmlElement("w:trackRevisions"))


def _mark_paragraph_inserted(paragraph: Paragraph, change_id: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    run_properties = paragraph_properties.find(qn("w:rPr"))
    if run_properties is None:
        run_properties = OxmlElement("w:rPr")
        paragraph_properties.append(run_properties)
    insertion = OxmlElement("w:ins")
    insertion.set(qn("w:id"), str(change_id))
    insertion.set(qn("w:author"), REVIEWER_AUTHOR)
    insertion.set(qn("w:date"), datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z"))
    run_properties.append(insertion)


def prepend_review_preface(document: Document, final_summary: PrefaceSummary) -> None:
    _insert_paragraph_at_start(document, "")
    feedback_paragraph = _insert_paragraph_at_start(
        document,
        final_summary.feedback_summary.strip(),
    )
    _ = feedback_paragraph
    text_summary_paragraph = _insert_paragraph_at_start(
        document,
        final_summary.global_feedback.strip(),
    )
    _ = text_summary_paragraph
    _insert_paragraph_at_start(
        document,
        final_summary.scientific_importance.strip(),
    )
    _insert_paragraph_at_start(
        document,
        final_summary.about_document.strip(),
    )
    _insert_paragraph_at_start(document, "Reviewer3")


def _insert_paragraph_at_start(document: Document, text: str | list[str]) -> Paragraph:
    if isinstance(text, list):
        text = "\n".join(f"• {item}" for item in text)

    if not document.paragraphs:
        paragraph = document.add_paragraph("")
    else:
        first_paragraph = document.paragraphs[0]._p
        new_paragraph_xml = OxmlElement("w:p")
        first_paragraph.addprevious(new_paragraph_xml)
        paragraph = Paragraph(new_paragraph_xml, first_paragraph.getparent())

    paragraph_change_id = _next_change_id(document)
    _mark_paragraph_inserted(paragraph, paragraph_change_id)
    if text:
        _append_tracked_insert(paragraph, text, change_id=paragraph_change_id + 1)
    return paragraph


def review_docx(input_docx: Path, output_docx: Path | None = None, reviewer: Any | None = None) -> Path:
    input_path = Path(input_docx)
    if not input_path.exists():
        raise FileNotFoundError(input_path)
    if input_path.suffix.lower() != ".docx":
        raise ValueError("Input file must be a .docx document.")

    requested_output = Path(output_docx) if output_docx else default_output_path(input_path)
    output_path = _non_overwriting_output_path(requested_output)
    reviewer = reviewer or LLMReviewer.from_env()

    doc = Document(str(input_path))
    _enable_track_revisions(doc)
    change_id = _next_change_id(doc)

    candidate_paragraphs = [paragraph for paragraph in doc.paragraphs if paragraph.text.strip()]
    document_text = "\n\n".join(paragraph.text for paragraph in candidate_paragraphs)

    feedback_entries: list[FeedbackEntry] = []
    if document_text and hasattr(reviewer, "review_document"):
        feedback_entries = reviewer.review_document(document_text=document_text)

    print(f"Collected {len(feedback_entries)} feedback entries.")
    # safe feedback to yml:
    #feedback_txt_path = output_path.with_suffix(".feedback.yml")
    #with open(feedback_txt_path, "w", encoding="utf-8") as f:
    #    yaml.dump([entry.__dict__ for entry in feedback_entries], f)

    counter = 0
    for entry in feedback_entries:
        if entry.feedback_type == "global_feedback":
            continue

        if not entry.word_group:
            continue

        paragraph = _find_paragraph_by_word_group(candidate_paragraphs, entry.word_group)
        if paragraph is None:
            continue

        if entry.feedback_type == "local_modification" and entry.suggestion:
            if not _preserves_layout_characters(entry.word_group, entry.suggestion):
                continue
            change_id_before = change_id
            change_id = add_tracked_suggestion(
                paragraph,
                entry.word_group,
                entry.suggestion,
                change_id=change_id,
            )
            if change_id != change_id_before:
                counter = counter + 1

        if entry.feedback_type == "regional_question" and entry.comment:
            add_comment_to_word_group(doc, paragraph, entry.word_group, entry.comment)

    if hasattr(reviewer, "compose_preface"):
        final_summary = reviewer.compose_preface(document_text=document_text, feedback_entries=feedback_entries)
    else:
        global_items = [item.comment for item in feedback_entries if item.feedback_type == "global_feedback" and item.comment]
        local_count = sum(1 for item in feedback_entries if item.feedback_type == "local_modification")
        regional_count = sum(1 for item in feedback_entries if item.feedback_type == "regional_question")
        final_summary = _fallback_preface_summary(
            fallback_global_feedback=global_items,
            fallback_local_count=local_count,
            fallback_regional_count=regional_count,
        )

    prepend_review_preface(doc, final_summary)

    doc.save(str(output_path))
    return output_path


def _find_paragraph_by_word_group(paragraphs: list[Any], word_group: str) -> Any | None:
    target = word_group.strip()
    if not target:
        return None

    target_lower = target.lower()
    for paragraph in paragraphs:
        paragraph_text = paragraph.text.strip()
        if paragraph_text and target_lower in paragraph_text.lower():
            return paragraph
    return None


def _preserves_layout_characters(original: str, suggestion: str) -> bool:
    protected = set("\r\n()[]{}<>")
    return [char for char in original if char in protected] == [
        char for char in suggestion if char in protected
    ]
